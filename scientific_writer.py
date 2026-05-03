import os
from docx import Document
import logging

class ScientificWriter:
    def __init__(self, drafts_path: str):
        self.drafts_path = drafts_path
        self.essay_formulas = {
            "introduction": "Hech kimga sir emaski, {topic} bugungi kunda dolzarb hisoblanadi. Ushbu maqolada biz {focus} masalalarini tahlil qilamiz.",
            "evidence": "Laboratoriya natijalarimiz shuni ko'rsatmoqdaki, {data_point}. Bu esa {thesis} g'oyasini tasdiqlaydi.",
            "conclusion": "Xulosa qilib aytganda, {summary}. Kelajakda bu yo'nalish {prospect} uchun katta poydevor bo'ladi."
        }

    def create_draft(self, title: str, content_blocks: dict):
        """Creates a professional .docx draft in the Drafts folder."""
        try:
            doc = Document()
            doc.add_heading(title, 0)

            # Using formulas to structure the text
            intro = self.essay_formulas["introduction"].format(
                topic=content_blocks.get("topic", "energetika tizimlari"),
                focus=content_blocks.get("focus", "samaradorlik")
            )
            doc.add_paragraph(intro)

            doc.add_heading("Asosiy qism va tahlil", level=1)
            body = content_blocks.get("body", "Bu yerda laboratoriya tahlillari va ilmiy dalillar keltiriladi.")
            doc.add_paragraph(body)

            evidence = self.essay_formulas["evidence"].format(
                data_point=content_blocks.get("data_point", "sinxron mashina quvvat koeffitsienti 0.85 ga teng"),
                thesis=content_blocks.get("thesis", "energiya yo'qotishlarini kamaytirish")
            )
            doc.add_paragraph(evidence)

            doc.add_heading("Xulosa", level=1)
            conclusion = self.essay_formulas["conclusion"].format(
                summary=content_blocks.get("summary", "gidravlika va energetika analogiyasi isbotlandi"),
                prospect=content_blocks.get("prospect", "yashil energetika")
            )
            doc.add_paragraph(conclusion)

            file_name = f"{title.replace(' ', '_')}.docx"
            file_path = os.path.join(self.drafts_path, file_name)
            doc.save(file_path)
            return file_path
        except Exception as e:
            logging.error(f"ScientificWriter Error: {e}")
            return None
